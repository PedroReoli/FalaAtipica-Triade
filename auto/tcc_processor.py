#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Automação para Processamento de TCC (DOCX)
==========================================

Este script processa arquivos de TCC em formato DOCX, extraindo texto e dividindo
em pedaços de 5000 linhas para análise.

Funcionalidades:
- Seleção de arquivo DOCX via interface gráfica
- Extração de texto do documento Word
- Divisão em pedaços de 5000 linhas
- Identificação de títulos vs texto normal
- Ignorar imagens e elementos gráficos
- Nomenclatura com linha de início e fim
"""

import tkinter as tk
from tkinter import filedialog, messagebox
import os
import re
from pathlib import Path
import sys

# Verificar se python-docx está instalado
try:
    from docx import Document
except ImportError:
    print("ERRO: Biblioteca python-docx não encontrada!")
    print("Instale com: pip install python-docx")
    sys.exit(1)

class TCCProcessor:
    def __init__(self):
        self.lines_per_chunk = 5000
        self.output_dir = "tcc_processed"
        
    def select_file(self):
        """Seleciona arquivo DOCX via interface gráfica"""
        root = tk.Tk()
        root.withdraw()  # Esconde a janela principal
        
        file_path = filedialog.askopenfilename(
            title="Selecione o arquivo do TCC (DOCX)",
            filetypes=[
                ("Arquivos Word", "*.docx"),
                ("Todos os arquivos", "*.*")
            ]
        )
        
        if not file_path:
            print("Nenhum arquivo selecionado.")
            return None
            
        return file_path
    
    def extract_text_from_docx(self, file_path):
        """
        Extrai texto de um arquivo DOCX
        """
        try:
            doc = Document(file_path)
            lines = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:  # Só adicionar linhas não vazias
                    lines.append(text)
            
            # Extrair texto de tabelas também
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            lines.append(text)
            
            return lines
            
        except Exception as e:
            print(f"Erro ao extrair texto do DOCX: {e}")
            return []
    
    def is_title(self, line):
        """
        Identifica se uma linha é um título baseado em padrões
        """
        line = line.strip()
        
        # Padrões para identificar títulos
        title_patterns = [
            r'^[A-Z][A-Z\s]+$',  # Tudo maiúsculo
            r'^\d+\.\s+[A-Z]',   # Numeração seguida de maiúscula
            r'^[A-Z][a-z]+(\s+[A-Z][a-z]+)*$',  # Primeira letra maiúscula de cada palavra
            r'^[A-Z][a-z]+(\s+[A-Z][a-z]+)*\s*$',  # Com espaços no final
            r'^\d+\.\d+\s+[A-Z]',  # Numeração decimal (1.1, 1.2, etc.)
            r'^[IVX]+\.\s+[A-Z]',  # Numeração romana
            r'^[A-Z][a-z]+(\s+[A-Z][a-z]+)*\s*:$',  # Título com dois pontos
            r'^APÊNDICE\s+[A-Z]',  # Apêndices
            r'^REFERÊNCIAS$',  # Referências
            r'^BIBLIOGRAFIA$',  # Bibliografia
        ]
        
        # Verificar se a linha é muito curta (provavelmente título)
        if len(line) < 100 and len(line) > 3:
            for pattern in title_patterns:
                if re.match(pattern, line):
                    return True
        
        return False
    
    def should_ignore_line(self, line):
        """
        Verifica se a linha deve ser ignorada (imagens, elementos gráficos, etc.)
        """
        line = line.strip()
        
        # Padrões para ignorar
        ignore_patterns = [
            r'^\[.*\]$',  # Linhas entre colchetes
            r'^\{.*\}$',  # Linhas entre chaves
            r'^<.*>$',    # Tags HTML/XML
            r'^[^\w\s]*$',  # Apenas símbolos especiais
            r'^[█▄▀▌▐░▒▓]+$',  # Caracteres de bloco (ASCII art)
            r'^[═║╔╗╚╝╠╣╦╩╬]+$',  # Caracteres de borda
            r'^[┌┐└┘├┤┬┴┼]+$',  # Caracteres de borda Unicode
            r'^\s*$',     # Linhas vazias ou apenas espaços
            r'^[^\w\s]*[█▄▀▌▐░▒▓][^\w\s]*$',  # Linhas com caracteres de bloco
            r'^[0-9]+$',  # Apenas números
            r'^[a-zA-Z0-9\s]*[0-9]{4}$',  # Referências bibliográficas
        ]
        
        for pattern in ignore_patterns:
            if re.match(pattern, line):
                return True
        
        return False
    
    def clean_line(self, line):
        """
        Limpa a linha removendo caracteres desnecessários
        """
        # Remove caracteres de controle exceto quebra de linha
        line = ''.join(char for char in line if ord(char) >= 32 or char == '\n')
        
        # Remove espaços extras no início e fim
        line = line.strip()
        
        # Remove múltiplos espaços
        line = re.sub(r'\s+', ' ', line)
        
        return line
    
    def process_file(self, file_path):
        """
        Processa o arquivo DOCX do TCC
        """
        try:
            # Criar diretório de saída
            Path(self.output_dir).mkdir(exist_ok=True)
            
            print(f"Extraindo texto do arquivo DOCX: {file_path}")
            
            # Extrair texto do DOCX
            lines = self.extract_text_from_docx(file_path)
            
            if not lines:
                print("Nenhum texto encontrado no arquivo DOCX!")
                return False
            
            print(f"Texto extraído: {len(lines)} linhas")
            
            # Filtrar e limpar linhas
            processed_lines = []
            for i, line in enumerate(lines, 1):
                if not self.should_ignore_line(line):
                    cleaned_line = self.clean_line(line)
                    if cleaned_line:  # Não adicionar linhas vazias
                        processed_lines.append({
                            'line_number': i,
                            'content': cleaned_line,
                            'is_title': self.is_title(cleaned_line)
                        })
            
            print(f"Linhas processadas: {len(processed_lines)}")
            
            if not processed_lines:
                print("Nenhuma linha válida encontrada após processamento!")
                return False
            
            # Dividir em chunks
            chunks = self.create_chunks(processed_lines)
            
            # Salvar chunks
            self.save_chunks(chunks, file_path)
            
            # Mostrar estatísticas
            self.show_statistics(chunks)
            
            print(f"Processamento concluído! {len(chunks)} arquivos criados em '{self.output_dir}'")
            
        except Exception as e:
            print(f"Erro ao processar arquivo: {e}")
            return False
        
        return True
    
    def create_chunks(self, processed_lines):
        """
        Divide as linhas processadas em chunks de 5000 linhas
        """
        chunks = []
        current_chunk = []
        
        for line_data in processed_lines:
            current_chunk.append(line_data)
            
            if len(current_chunk) >= self.lines_per_chunk:
                chunks.append(current_chunk)
                current_chunk = []
        
        # Adicionar o último chunk se não estiver vazio
        if current_chunk:
            chunks.append(current_chunk)
        
        return chunks
    
    def save_chunks(self, chunks, original_file_path):
        """
        Salva os chunks em arquivos separados
        """
        base_name = Path(original_file_path).stem
        
        for i, chunk in enumerate(chunks, 1):
            if not chunk:
                continue
                
            start_line = chunk[0]['line_number']
            end_line = chunk[-1]['line_number']
            
            # Nome do arquivo com linha de início e fim
            filename = f"{base_name}_Parte_{i:02d}_Linhas_{start_line:05d}-{end_line:05d}.txt"
            filepath = Path(self.output_dir) / filename
            
            with open(filepath, 'w', encoding='utf-8') as file:
                file.write(f"TCC Processado - Parte {i}\n")
                file.write(f"Arquivo original: {Path(original_file_path).name}\n")
                file.write(f"Linhas do arquivo original: {start_line} - {end_line}\n")
                file.write(f"Total de linhas neste arquivo: {len(chunk)}\n")
                file.write("=" * 80 + "\n\n")
                
                for line_data in chunk:
                    line_num = line_data['line_number']
                    content = line_data['content']
                    is_title = line_data['is_title']
                    
                    # Marcar títulos
                    if is_title:
                        file.write(f"[TÍTULO - Linha {line_num}] {content}\n")
                    else:
                        file.write(f"[Linha {line_num}] {content}\n")
            
            print(f"Arquivo criado: {filename}")
    
    def show_statistics(self, chunks):
        """
        Mostra estatísticas do processamento
        """
        total_lines = sum(len(chunk) for chunk in chunks)
        total_titles = sum(
            sum(1 for line in chunk if line['is_title'])
            for chunk in chunks
        )
        
        print("\n" + "=" * 50)
        print("ESTATÍSTICAS DO PROCESSAMENTO")
        print("=" * 50)
        print(f"Total de arquivos criados: {len(chunks)}")
        print(f"Total de linhas processadas: {total_lines}")
        print(f"Total de títulos identificados: {total_titles}")
        print(f"Linhas por arquivo: {self.lines_per_chunk}")
        print("=" * 50)

def main():
    """
    Função principal
    """
    print("=" * 60)
    print("PROCESSADOR DE TCC DOCX - FalaAtípica")
    print("=" * 60)
    print("Este script processa arquivos de TCC em formato DOCX, extraindo")
    print("texto e dividindo em pedaços de 5000 linhas para análise.")
    print()
    
    processor = TCCProcessor()
    
    # Selecionar arquivo
    file_path = processor.select_file()
    if not file_path:
        return
    
    print(f"Arquivo selecionado: {file_path}")
    print("Iniciando processamento...")
    print()
    
    # Processar arquivo
    success = processor.process_file(file_path)
    
    if success:
        print("\nProcessamento concluído com sucesso!")
        print(f"Arquivos salvos em: {processor.output_dir}")
        
        # Perguntar se quer abrir a pasta
        root = tk.Tk()
        root.withdraw()
        
        if messagebox.askyesno("Sucesso", "Processamento concluído! Deseja abrir a pasta com os arquivos?"):
            os.startfile(processor.output_dir)
    else:
        print("\nErro no processamento!")
        messagebox.showerror("Erro", "Ocorreu um erro durante o processamento!")

if __name__ == "__main__":
    main()
